"""
Document parsing service for extracting text from various file formats.

Supports: .txt, .pdf, .docx files
"""

import io
from typing import BinaryIO, Union
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


class DocumentParsingError(Exception):
    """Raised when document parsing fails."""
    pass


class UnsupportedFileTypeError(DocumentParsingError):
    """Raised when file type is not supported."""
    pass


class DocumentParser:
    """
    Service for parsing documents and extracting text content.

    Supports multiple file formats:
    - Plain text (.txt)
    - PDF (.pdf) - requires PyPDF2
    - Microsoft Word (.docx) - requires python-docx
    """

    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx'}

    @classmethod
    def parse_document(cls, file: Union[BinaryIO, bytes], file_type: str) -> str:
        """
        Parse a document and extract its text content.

        Args:
            file: File object or bytes to parse
            file_type: File extension (e.g., 'txt', 'pdf', 'docx')

        Returns:
            Extracted text content as a string

        Raises:
            UnsupportedFileTypeError: If file type is not supported
            DocumentParsingError: If parsing fails
        """
        # Normalize file type
        file_type = file_type.lower().lstrip('.')

        if f'.{file_type}' not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"File type '.{file_type}' is not supported. "
                f"Supported types: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        # Convert bytes to file-like object if needed
        if isinstance(file, bytes):
            file = io.BytesIO(file)

        # Route to appropriate parser
        try:
            if file_type == 'txt':
                return cls.extract_text_from_txt(file)
            elif file_type == 'pdf':
                return cls.extract_text_from_pdf(file)
            elif file_type == 'docx':
                return cls.extract_text_from_docx(file)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
        except DocumentParsingError:
            raise
        except Exception as e:
            raise DocumentParsingError(f"Failed to parse {file_type} document: {str(e)}") from e

    @staticmethod
    def extract_text_from_txt(file: BinaryIO) -> str:
        """
        Extract text from a plain text file.

        Args:
            file: File object to read

        Returns:
            Text content

        Raises:
            DocumentParsingError: If reading fails
        """
        try:
            # Try UTF-8 first, fall back to latin-1 if that fails
            content = file.read()
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
        except Exception as e:
            raise DocumentParsingError(f"Failed to read text file: {str(e)}") from e

    @staticmethod
    def extract_text_from_pdf(file: BinaryIO) -> str:
        """
        Extract text from a PDF file.

        Args:
            file: File object to read

        Returns:
            Extracted text content

        Raises:
            DocumentParsingError: If PDF reading fails or PyPDF2 is not installed
        """
        if PdfReader is None:
            raise DocumentParsingError(
                "PyPDF2 is not installed. Install it with: pip install PyPDF2"
            )

        try:
            reader = PdfReader(file)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            if not text_parts:
                return ""

            return "\n\n".join(text_parts)
        except Exception as e:
            raise DocumentParsingError(f"Failed to read PDF file: {str(e)}") from e

    @staticmethod
    def extract_text_from_docx(file: BinaryIO) -> str:
        """
        Extract text from a Microsoft Word (.docx) file.

        Args:
            file: File object to read

        Returns:
            Extracted text content

        Raises:
            DocumentParsingError: If DOCX reading fails or python-docx is not installed
        """
        if Document is None:
            raise DocumentParsingError(
                "python-docx is not installed. Install it with: pip install python-docx"
            )

        try:
            doc = Document(file)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            if not text_parts:
                return ""

            return "\n\n".join(text_parts)
        except Exception as e:
            raise DocumentParsingError(f"Failed to read DOCX file: {str(e)}") from e

    @classmethod
    def parse_file_path(cls, file_path: Union[str, Path]) -> str:
        """
        Convenience method to parse a document from a file path.

        Args:
            file_path: Path to the file to parse

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            UnsupportedFileTypeError: If file type is not supported
            DocumentParsingError: If parsing fails
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = path.suffix.lstrip('.')

        with open(path, 'rb') as f:
            return cls.parse_document(f, file_type)
